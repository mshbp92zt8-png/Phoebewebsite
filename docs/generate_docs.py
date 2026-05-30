"""
Generates Privacy Policy and Terms of Use .docx files for BodyMindOT.
Run: python3 generate_docs.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

GOLD = RGBColor(0xC4, 0xA8, 0x82)
CHARCOAL = RGBColor(0x1C, 0x18, 0x14)
TEXT = RGBColor(0x2A, 0x2A, 0x2A)
TEXT_LIGHT = RGBColor(0x6B, 0x65, 0x60)


def set_font(run, size=11, bold=False, color=None, italic=False):
    run.font.name = "Georgia"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    if level == 1:
        set_font(run, size=16, bold=False, color=CHARCOAL)
        run.font.name = "Georgia"
    else:
        set_font(run, size=13, bold=False, color=CHARCOAL)
    return p


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        set_font(r, bold=True, color=TEXT)
    run = p.add_run(text)
    set_font(run, color=TEXT)
    return p


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        set_font(r, bold=True, color=TEXT)
    run = p.add_run(text)
    set_font(run, color=TEXT)
    return p


def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "EDE8E0")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_label(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text.upper())
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    run.font.color.rgb = GOLD
    run.font.bold = True


def setup_doc(title):
    doc = Document()
    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(3)
    # Title block
    add_label(doc, "BodyMindOT · Legal")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(title)
    run.font.name = "Georgia"
    run.font.size = Pt(26)
    run.font.color.rgb = CHARCOAL
    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(16)
    r = meta.add_run("Last updated: May 2026  ·  Effective date: 1 June 2026")
    r.font.name = "Calibri"
    r.font.size = Pt(9)
    r.font.color.rgb = TEXT_LIGHT
    add_divider(doc)
    return doc


# ─── PRIVACY POLICY ──────────────────────────────────────────────────────────

def build_privacy_policy():
    doc = setup_doc("Privacy Policy")

    intro = doc.add_paragraph()
    intro.paragraph_format.space_after = Pt(12)
    r = intro.add_run(
        "This policy explains what personal data BodyMindOT collects, why, and how it is used. "
        "It is written in accordance with the UK General Data Protection Regulation (UK GDPR) "
        "and the Data Protection Act 2018."
    )
    set_font(r, color=TEXT)

    # 1
    add_heading(doc, "1. Who We Are")
    add_body(doc, (
        "BodyMindOT is the trading name of Phoebe Hsieh, a sole practitioner and HCPC-registered "
        "Occupational Therapist providing online mental health and wellbeing consultations."
    ))
    add_body(doc, "Phoebe Hsieh (trading as BodyMindOT)", bold_prefix="Data Controller: ")
    add_body(doc, "phoebehsieh@bodymindot.com", bold_prefix="Contact: ")
    add_divider(doc)

    # 2
    add_heading(doc, "2. What Personal Data We Collect")
    add_heading(doc, "Contact and enquiry data", level=2)
    for item in ["Name", "Email address", "Any information you include in the message field of our contact form"]:
        add_bullet(doc, item)
    add_heading(doc, "Health data (if voluntarily provided)", level=2)
    add_body(doc, (
        "If you choose to describe health, mental health, or wellbeing concerns in your message, "
        "this constitutes special category data under UK GDPR and is handled with heightened care. "
        "We do not ask you to submit health information through the website — any such information "
        "is provided entirely at your discretion."
    ))
    add_heading(doc, "Technical data", level=2)
    for item in [
        "IP address (collected automatically by our hosting provider)",
        "Browser type and version",
        "Pages visited and time spent (if analytics are active)",
    ]:
        add_bullet(doc, item)
    add_divider(doc)

    # 3
    add_heading(doc, "3. How We Collect Your Data")
    add_bullet(doc, "when you submit an enquiry via the website", bold_prefix="Contact form: ")
    add_bullet(doc, "when you contact us directly at our email address", bold_prefix="Email: ")
    add_bullet(doc, "technical data collected by our web hosting provider when you visit the site", bold_prefix="Automatically: ")
    add_divider(doc)

    # 4
    add_heading(doc, "4. Legal Basis for Processing")
    add_bullet(doc, "to respond to enquiries you have sent to us", bold_prefix="Legitimate interests (Article 6(1)(f)): ")
    add_bullet(doc, "where processing is necessary to provide occupational therapy services you have requested", bold_prefix="Contract (Article 6(1)(b)): ")
    add_bullet(doc, "for any optional communications or marketing, where explicitly obtained", bold_prefix="Consent (Article 6(1)(a)): ")
    add_bullet(doc, "where we are required by law to retain records", bold_prefix="Legal obligation (Article 6(1)(c)): ")
    add_body(doc, (
        "For special category health data, the lawful basis is explicit consent (Article 9(2)(a)) "
        "or health and social care purposes (Article 9(2)(h)) where applicable."
    ))
    add_divider(doc)

    # 5
    add_heading(doc, "5. How We Use Your Data")
    for item in [
        "To respond to your enquiry or consultation request",
        "To provide and manage occupational therapy services",
        "To maintain professional records as required by HCPC standards of conduct",
        "To improve the website and user experience (using aggregated, anonymised analytics)",
    ]:
        add_bullet(doc, item)
    add_body(doc, "We will never sell your personal data, or share it with third parties for their own marketing purposes.")
    add_divider(doc)

    # 6
    add_heading(doc, "6. Data Retention")
    add_bullet(doc, "deleted within 12 months if no therapeutic relationship is established", bold_prefix="Enquiry data (non-client): ")
    add_bullet(doc, "retained for a minimum of 7 years after the end of the therapeutic relationship, in line with HCPC guidance", bold_prefix="Client records: ")
    add_bullet(doc, "retained in accordance with our hosting provider's policies (typically up to 26 months)", bold_prefix="Technical / analytics data: ")
    add_bullet(doc, "retained for up to 12 months for record-keeping purposes", bold_prefix="Workshop attendance lists: ")
    add_divider(doc)

    # 7
    add_heading(doc, "7. Third Parties and Data Sharing")
    add_body(doc, "We do not sell or rent your personal data. We may share data with the following categories of third parties only where necessary:")
    add_bullet(doc, "to send and receive communications (e.g. Google Workspace)", bold_prefix="Email service provider: ")
    add_bullet(doc, "who may process technical data as part of hosting the site", bold_prefix="Website hosting provider: ")
    add_bullet(doc, (
        "used for one-to-one online consultations. Call data is processed by Google LLC under "
        "Google's privacy policy (policies.google.com/privacy). No session recordings are made "
        "without your explicit prior consent."
    ), bold_prefix="Google Meet: ")
    add_bullet(doc, "in the event of a professional fitness-to-practise investigation, we may be required to share records with the HCPC", bold_prefix="Regulatory bodies: ")
    add_bullet(doc, "where required by law, court order, or to protect the safety of you or others", bold_prefix="Legal or statutory obligation: ")
    add_body(doc, "All third-party processors are required to handle your data securely and in accordance with UK data protection law.")
    add_divider(doc)

    # 8
    add_heading(doc, "8. Workshops")
    add_body(doc, "BodyMindOT runs in-person group workshops. If you register for a workshop, we may collect:")
    for item in [
        "Name and email address for registration and joining instructions",
        "Emergency contact details where required for the venue",
        "Any access or support needs you choose to disclose",
    ]:
        add_bullet(doc, item)
    add_body(doc, "This data is used solely to administer the workshop and is not shared with other attendees.")
    add_body(doc, (
        "Please note: workshops are group settings. You should not share personal health information "
        "during a workshop that you would not be comfortable others in the group hearing. BodyMindOT "
        "is not able to guarantee the confidentiality of any disclosures made by participants in a "
        "group environment."
    ))
    add_divider(doc)

    # 9
    add_heading(doc, "9. International Data Transfers")
    add_body(doc, (
        "Where your data is transferred outside the UK (for example, by a third-party service provider), "
        "we ensure that appropriate safeguards are in place — such as UK adequacy regulations or Standard "
        "Contractual Clauses — to protect your data to an equivalent standard."
    ))
    add_divider(doc)

    # 10
    add_heading(doc, "10. Cookies")
    add_body(doc, (
        "This website uses essential cookies only — small files placed on your device that are strictly "
        "necessary for the site to function. We do not use advertising cookies or cross-site tracking cookies."
    ))
    add_body(doc, (
        "If we introduce any non-essential cookies in future (such as analytics), we will update this "
        "policy and seek your consent via the cookie banner."
    ))
    add_body(doc, "You can control cookies through your browser settings at any time. Disabling essential cookies may affect site functionality.")
    add_divider(doc)

    # 11
    add_heading(doc, "11. Your Rights Under UK GDPR")
    add_bullet(doc, "to obtain a copy of the personal data we hold about you", bold_prefix="Right of access: ")
    add_bullet(doc, "to have inaccurate data corrected", bold_prefix="Right to rectification: ")
    add_bullet(doc, "to request deletion of your data, subject to legal retention obligations", bold_prefix="Right to erasure: ")
    add_bullet(doc, "to limit how we use your data in certain circumstances", bold_prefix="Right to restrict processing: ")
    add_bullet(doc, "to receive your data in a structured, machine-readable format", bold_prefix="Right to data portability: ")
    add_bullet(doc, "to processing based on legitimate interests", bold_prefix="Right to object: ")
    add_bullet(doc, "where processing is based on consent, you may withdraw it at any time without affecting the lawfulness of prior processing", bold_prefix="Right to withdraw consent: ")
    add_body(doc, "To exercise any of these rights, contact us at phoebehsieh@bodymindot.com. We will respond within one calendar month.")
    add_divider(doc)

    # 12
    add_heading(doc, "12. How to Complain")
    add_body(doc, (
        "If you believe your data has been handled incorrectly, please contact us in the first instance. "
        "You also have the right to lodge a complaint with the Information Commissioner's Office (ICO):"
    ))
    add_bullet(doc, "ico.org.uk", bold_prefix="Website: ")
    add_bullet(doc, "0303 123 1113", bold_prefix="Helpline: ")
    add_divider(doc)

    # 13
    add_heading(doc, "13. Changes to This Policy")
    add_body(doc, (
        "We may update this Privacy Policy from time to time. The 'Last updated' date at the top of "
        "this page will reflect any changes. Continued use of the website after changes are posted "
        "constitutes acceptance of the updated policy."
    ))
    add_divider(doc)

    # 14
    add_heading(doc, "14. Contact")
    add_body(doc, "Phoebe Hsieh — BodyMindOT")
    add_body(doc, "phoebehsieh@bodymindot.com", bold_prefix="Email: ")

    doc.save("privacy-policy.docx")
    print("✓ privacy-policy.docx created")


# ─── TERMS OF USE ─────────────────────────────────────────────────────────────

def build_terms_of_use():
    doc = setup_doc("Terms of Use")

    intro = doc.add_paragraph()
    intro.paragraph_format.space_after = Pt(12)
    r = intro.add_run(
        "By accessing or using this website, you agree to be bound by these Terms of Use. "
        "Please read them carefully. If you do not agree, please do not use this website."
    )
    set_font(r, color=TEXT)

    # 1
    add_heading(doc, "1. About This Website")
    add_body(doc, (
        "This website is operated by Phoebe Hsieh, trading as BodyMindOT, an HCPC-registered "
        "Occupational Therapist. The website provides information about occupational therapy "
        "services and facilitates initial enquiries."
    ))
    add_body(doc, "phoebehsieh@bodymindot.com", bold_prefix="Contact: ")
    add_divider(doc)

    # 2
    add_heading(doc, "2. Healthcare Disclaimer")
    add_body(doc, (
        "The content on this website is provided for general informational purposes only. It does "
        "not constitute medical advice, clinical assessment, diagnosis, or treatment, and must not "
        "be relied upon as a substitute for professional healthcare advice."
    ))
    add_body(doc, (
        "If you have concerns about your physical or mental health, please consult a qualified "
        "healthcare professional or, in an emergency, contact your local emergency services or "
        "the NHS (999 / 111)."
    ))
    add_body(doc, (
        "One-to-one online consultations are conducted via Google Meet. By participating, you also "
        "agree to Google's Terms of Service (policies.google.com/terms). Sessions are not recorded "
        "without your explicit prior consent."
    ))
    add_body(doc, (
        "In-person group workshops take place at a physical venue. Workshops are educational in "
        "nature and do not constitute individual clinical treatment or assessment. Participants should "
        "not disclose personal health information they would not be comfortable sharing with the group."
    ))
    add_body(doc, (
        "A therapeutic relationship with BodyMindOT begins only once you have received written "
        "confirmation of your first individual appointment. Prior to that point, no professional "
        "duty of care exists through this website."
    ))
    add_divider(doc)

    # 3
    add_heading(doc, "3. Acceptable Use")
    add_body(doc, "You may use this website for lawful purposes only. You must not:")
    for item in [
        "Use the website in any way that violates applicable UK or international laws or regulations",
        "Transmit any unsolicited or unauthorised advertising or promotional material",
        "Attempt to gain unauthorised access to any part of the website or its infrastructure",
        "Submit false, misleading, or offensive content through any contact form",
        "Engage in any conduct that could damage, disable, or impair the website",
    ]:
        add_bullet(doc, item)
    add_divider(doc)

    # 4
    add_heading(doc, "4. Intellectual Property")
    add_body(doc, (
        "All content on this website — including text, images, logos, graphics, and design — is the "
        "property of BodyMindOT or its licensors and is protected by copyright and other intellectual "
        "property laws."
    ))
    add_body(doc, (
        "You may view and print content for your own personal, non-commercial use. You must not "
        "reproduce, distribute, modify, or commercially exploit any part of this website without "
        "our prior written consent."
    ))
    add_divider(doc)

    # 5
    add_heading(doc, "5. Accuracy of Information")
    add_body(doc, (
        "We take reasonable care to ensure the information on this website is accurate and up to date. "
        "However, we make no warranty that the content is complete, current, or error-free. We reserve "
        "the right to update or correct information at any time without notice."
    ))
    add_divider(doc)

    # 6
    add_heading(doc, "6. Links to Third-Party Websites")
    add_body(doc, (
        "This website may contain links to external websites for your convenience. These links do not "
        "signify our endorsement of those sites. We have no control over the content or availability "
        "of external sites and accept no responsibility for them or for any loss or damage that may "
        "arise from your use of them."
    ))
    add_divider(doc)

    # 7
    add_heading(doc, "7. Limitation of Liability")
    add_body(doc, "To the fullest extent permitted by law, BodyMindOT excludes all liability for any loss or damage — whether direct, indirect, or consequential — arising from:")
    for item in [
        "Your use of, or inability to use, this website",
        "Reliance on any information contained on this website",
        "Any interruption, suspension, or discontinuation of the website",
    ]:
        add_bullet(doc, item)
    add_body(doc, (
        "Nothing in these Terms limits our liability for death or personal injury caused by negligence, "
        "fraud, or any other matter where exclusion of liability is not permitted by law."
    ))
    add_divider(doc)

    # 8
    add_heading(doc, "8. Privacy")
    add_body(doc, (
        "Your use of this website is also governed by our Privacy Policy, which is incorporated into "
        "these Terms by reference. By using this website, you consent to the data practices described "
        "in that policy."
    ))
    add_divider(doc)

    # 9
    add_heading(doc, "9. Governing Law")
    add_body(doc, (
        "These Terms of Use are governed by and construed in accordance with the laws of England and "
        "Wales. Any disputes arising from your use of this website shall be subject to the exclusive "
        "jurisdiction of the courts of England and Wales."
    ))
    add_divider(doc)

    # 10
    add_heading(doc, "10. Changes to These Terms")
    add_body(doc, (
        "We reserve the right to amend these Terms of Use at any time. The 'Last updated' date at the "
        "top of this page will reflect any changes. Continued use of the website after updates are "
        "posted constitutes your acceptance of the revised Terms."
    ))
    add_divider(doc)

    # 11
    add_heading(doc, "11. Contact")
    add_body(doc, "Phoebe Hsieh — BodyMindOT")
    add_body(doc, "phoebehsieh@bodymindot.com", bold_prefix="Email: ")

    doc.save("terms-of-use.docx")
    print("✓ terms-of-use.docx created")


if __name__ == "__main__":
    build_privacy_policy()
    build_terms_of_use()
